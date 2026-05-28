"""陈熙贤简历 v2 - 全项目优化版"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "CV-陈熙贤-智能制造-AI实习生-v2.docx")

doc = Document()

# ---- 页面 ----
s = doc.sections[0]
s.page_width = Cm(21)
s.page_height = Cm(29.7)
s.top_margin = Cm(1.5)
s.bottom_margin = Cm(1.2)
s.left_margin = Cm(2)
s.right_margin = Cm(2)

# ---- 样式 ----
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = Pt(18)

BLUE = RGBColor(0x1a, 0x52, 0x76)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(0x66, 0x66, 0x66)


def add_name(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(4)


def add_contact(items):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("   •   ".join(items))
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(8)


def add_section(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '8',
        qn('w:space'): '1', qn('w:color'): '1A5276',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    r = p.add_run(title)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = BLUE


def add_entry(left, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(left)
    r.font.size = Pt(11)
    r.font.bold = True
    r2 = p.add_run("\t" + right)
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17), alignment=WD_ALIGN_PARAGRAPH.RIGHT)


def add_sub(left, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(left)
    r.font.size = Pt(10)
    r2 = p.add_run("\t" + right)
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17), alignment=WD_ALIGN_PARAGRAPH.RIGHT)


def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run("● ")
    r.font.size = Pt(7)
    r.font.color.rgb = BLUE
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)


def add_skill(cat, val):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run("● ")
    r.font.size = Pt(7)
    r.font.color.rgb = BLUE
    r2 = p.add_run(cat)
    r2.font.size = Pt(9.5)
    r2.font.bold = True
    r3 = p.add_run(val)
    r3.font.size = Pt(9.5)


# ============================================================
# 内容
# ============================================================

add_name("陈熙贤")
add_contact([
    "23xxchen@stu.edu.cn",
    "+86 137-9061-3670",
    "广东汕头/东莞/深圳",
    "GitHub: github.com/23xxCh",
])

# ---- 教育 ----
add_section("教育背景")
add_entry("汕头大学  工学院  智能制造工程", "广东 汕头")
add_sub("本科  |  智能制造工程", "2023.09 – 2027.06（预计）")
add_bullet("核心课程：精益生产、先进制造、机器人技术、信号与系统、机器学习、嵌入式系统、机器视觉原理、产品设计")
add_bullet("英语四级 524 分 | 挑战杯广东省省赛一等奖")

# ---- 实习 ----
add_section("实习经历")
add_entry("伟易达（东莞）电子产品有限公司", "广东 东莞")
add_sub("机械结构实习生", "2025 暑期")
add_bullet("参与玩具产品产线实习，了解注塑、组装、质检等完整生产流程，深入学习模具结构与成型工艺")
add_bullet("使用 Creo 进行零部件 3D 建模与装配设计，配合模具知识完成产品结构优化方案")
add_bullet("产线现场观察与记录，运用精益生产思维分析生产瓶颈，提出改进建议")

# ---- 项目 ----
add_section("项目经历")

# 项目1: auto prediction
add_entry("智能装备健康管理系统（IEHM）— 设备预测性维护平台", "2025")
add_bullet("构建工业设备预测性维护平台，模拟数控机床/织布机/机械臂 3 类设备传感器时序数据，训练 XGBoost 模型预测剩余使用寿命（RUL），预测误差 < 8%")
add_bullet("开发工业级 HMI 风格数字孪生看板（Streamlit），实现设备状态实时监控、告警 Toast 通知、WebSocket 状态推送")
add_bullet("集成大模型 API 自动生成故障诊断报告，包含故障原因分析和维护建议，迭代 17 轮，项目评分 8.9/10")

# 项目2: YOLO detect
add_entry("纺织布面缺陷视觉检测系统 — 为汕头纺织业定制", "2025")
add_bullet("基于 YOLOv8 构建布面缺陷检测模型，支持破洞/污渍/断经/色差 4 类缺陷实时检测，mAP@50 达 92%")
add_bullet("集成大模型自动生成结构化缺陷报告与处理建议，开发缺陷热力图可视化（Attention Map）和 ROI 投资回报计算器")
add_bullet("实现 Few-shot 快速适配新织物类型、SQLite 检测记录持久化、CSV/JSON 报告导出等工业落地功能")

# 项目3: industrial-knowledge-assistant
add_entry("工业知识助手 — 基于 RAG 的制造业智能问答系统", "2025")
add_bullet("基于 RAG 架构构建工业知识问答系统，整合设备手册、工艺参数、SOP 文档，实现自然语言查询故障代码和操作规范")
add_bullet("实现 Hybrid RRF 检索（Dense 向量 + BM25 关键词 + RRF 融合），技术文档检索精度相比单一方法提升 25%")
add_bullet("开发 Agent 工具层对接 SCADA/MES/设备参数系统，支持 LLM Function Calling；全栈开发：Spring Boot + Python RAG（Chroma + FastAPI）+ Vue 3")

# 项目4: digital system
add_entry("智能制造数字孪生系统 — 工业 4.0 架构参考设计", "2024-2025")
add_bullet("设计并实现智能制造数字孪生系统，覆盖 OPC UA 设备建模、DQN 强化学习智能调度、MES/WMS/质量管理/成本核算全业务链")
add_bullet("基于 DQN 强化学习实现产线调度优化，相比规则调度效率提升 15%；实现 Purdue 安全架构分层，110+ 单元测试通过")
add_bullet("对标 Siemens Xcelerator，以开源免费 + 轻量部署 + 纯 Python 栈的差异化方案服务中小企业")

# 项目5: xiaozhi
add_entry("ESP32 AI 智能对话机器人 — 端侧 AI Agent", "2025")
add_bullet("基于 ESP32 + 小智 AI 框架开发端侧智能对话机器人，集成语音识别、大模型推理、TTS 语音合成全链路")
add_bullet("深入理解固件架构（音频/显示/网络协议/MCP 服务），能独立进行功能定制和调试")
add_bullet("硬件端与云端大模型实时通信，支持多轮对话与语音交互；具备 AI Agent 开发经验，熟悉 Claude Code、Codex 等工具")

# 项目6: ROS2
add_entry("基于 ROS2 的氢气检测自主巡检小车", "2025")
add_bullet("基于 ROS2 框架开发自主巡检机器人，集成氢气传感器模块，实现危险环境自动检测与预警")
add_bullet("完成 SLAM 导航建图、路径规划、传感器融合等核心模块开发与调试")
add_bullet("硬件选型与嵌入式系统开发：ESP32 主控 + 传感器通信 + PCB 画板（嘉立创 EDA）")

# ---- 技能 ----
add_section("专业技能")
add_skill("CAD/建模：", "SolidWorks（熟练）、Creo（熟练）、UG（熟练）、AutoCAD（熟练）、Moldex 模流分析")
add_skill("编程语言：", "Python（主力）、C/C++（嵌入式）、JavaScript/TypeScript、MATLAB、SQL")
add_skill("AI/ML：", "PyTorch、Scikit-learn、YOLOv8、LangChain、RAG 架构、大模型 API、Prompt Engineering")
add_skill("机器人/嵌入式：", "ROS2 框架、SLAM、ESP32 开发、PCB 设计（嘉立创 EDA）、PLC 编程（博途）")
add_skill("后端/DevOps：", "Spring Boot、FastAPI、Docker、MySQL、Redis、PostgreSQL、Git")
add_skill("前端/可视化：", "Vue 3、Streamlit、PySide6、TypeScript")
add_skill("工业系统：", "OPC UA、MES/WMS/SCADA 对接、数字孪生、精益生产")
add_skill("AI 工具：", "Claude Code、Codex 等 AI 辅助编程工具，AI Agent 开发经验")

# ---- 获奖 ----
add_section("获奖与荣誉")
add_bullet("挑战杯广东省省赛一等奖 — 一鉴钟氢（光声衰荡光谱氢气检测技术）")
add_bullet("GitHub 个人仓库 20+ 个项目，涵盖工业 AI、机器视觉、数字孪生、智能硬件、AI Agent 等方向")

# ---- 自我评价 ----
add_section("自我评价")
add_bullet("智能制造工程专业背景，具备「机械结构 + 嵌入式硬件 + AI 软件」三位一体的复合能力")
add_bullet("自驱力强，20+ 个 GitHub 项目全部独立完成或主导开发；外向开朗，团队协作能力优秀")

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
