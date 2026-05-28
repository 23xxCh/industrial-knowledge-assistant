"""陈熙贤简历 - docx 格式"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "CV-陈熙贤-智能制造工程师.docx")

doc = Document()

# ---- 页面设置 ----
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# ---- 样式 ----
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = Pt(18)

BLUE = RGBColor(0x1a, 0x52, 0x76)
BLACK = RGBColor(0, 0, 0)
GRAY = 0x666666


def add_name(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(4)


def add_contact(items):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("   •   ".join(items))
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.paragraph_format.space_after = Pt(10)


def add_section(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    # 底部边框作为分隔线
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '8',
        qn('w:space'): '1',
        qn('w:color'): '1A5276',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    r = p.add_run(title)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = BLUE


def add_entry(left, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    # 左侧
    r = p.add_run(left)
    r.font.size = Pt(11)
    r.font.bold = True
    # 右侧（用 tab）
    r2 = p.add_run("\t" + right)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    # 设置 tab stop 到右对齐
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17), alignment=WD_ALIGN_PARAGRAPH.RIGHT)


def add_sub(left, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(left)
    r.font.size = Pt(10)
    r2 = p.add_run("\t" + right)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17), alignment=WD_ALIGN_PARAGRAPH.RIGHT)


def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run("● ")
    r.font.size = Pt(8)
    r.font.color.rgb = BLUE
    r2 = p.add_run(text)
    r2.font.size = Pt(10)


def add_skill(cat, val):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run("● ")
    r.font.size = Pt(8)
    r.font.color.rgb = BLUE
    r2 = p.add_run(cat)
    r2.font.size = Pt(10)
    r2.font.bold = True
    r3 = p.add_run(val)
    r3.font.size = Pt(10)


# ============================================================
# 内容
# ============================================================

add_name("陈熙贤")
add_contact([
    "23xxchen@stu.edu.cn",
    "+86 137-9061-3670",
    "广东汕头",
    "GitHub: github.com/23xxCh",
])

# ---- 教育背景 ----
add_section("教育背景")
add_entry("汕头大学  工学院  智能制造工程", "广东 汕头")
add_sub("本科  |  智能制造工程", "2023.09 – 2027.06（预计）")
add_bullet("核心课程：精益生产、先进制造、机器人技术、信号与系统、产品设计、机器学习、嵌入式系统、机器视觉原理")
add_bullet("英语四级 524 分，具备英文技术文档阅读能力")

# ---- 实习经历 ----
add_section("实习经历")
add_entry("伟易达（东莞）电子产品有限公司", "广东 东莞")
add_sub("机械结构实习生", "2025 暑期")
add_bullet("参与玩具产品产线实习，了解注塑、组装、质检等完整生产流程，深入学习模具结构与成型工艺。")
add_bullet("使用 Creo 进行零部件 3D 建模与装配设计，配合模具知识完成产品结构优化方案。")
add_bullet("产线现场观察与记录，协助分析生产瓶颈，提出基于精益生产的改进建议。")

# ---- 项目经历 ----
add_section("项目经历")

add_entry("一鉴钟氢 — 基于光声衰荡光谱的氢气检测技术", "2025")
add_bullet("挑战杯广东省省赛一等奖项目。负责光声衰荡光谱检测系统的硬件搭建与信号处理。")
add_bullet("使用 MATLAB 进行光谱信号采集与数据分析，优化检测精度，实现 ppb 级氢气浓度检测。")
add_bullet("撰写技术文档与实验报告，参与路演答辩，从 200+ 支队伍中脱颖而出。")

add_entry("基于 ROS2 的氢气检测自主巡检小车", "2025")
add_bullet("基于 ROS2 框架开发自主巡检机器人，集成氢气传感器模块，实现危险环境自动检测与预警。")
add_bullet("完成导航建图（SLAM）、路径规划、传感器融合等核心模块开发与调试。")
add_bullet("硬件选型与嵌入式系统开发，ESP32 主控 + 传感器通信，PCB 画板（嘉立创 EDA）。")

add_entry("AI Agent 智能对话机器人", "2025")
add_bullet("基于小智 AI 与 ESP-Claw 框架开发智能对话机器人，集成语音识别与大模型推理能力。")
add_bullet("实现硬件端（ESP32）与云端大模型的实时通信，支持多轮对话与语音交互。")
add_bullet("具备 AI Agent 开发经验，熟悉 Claude Code、Codex 等 AI 辅助编程工具。")

add_entry("数据分析与预测建模", "2024")
add_bullet("对汽车销售数据进行清洗、特征工程与可视化分析，使用 Python 构建回归预测模型。")
add_bullet("基于 MATLAB 完成桥梁静力学有限元分析（伏图软件），输出结构应力分布报告。")
add_bullet("使用嘉立创 EDA 完成 PCB 电路设计与打样，了解博途软件并完成 PLC 程序编写。")

# ---- 技能 ----
add_section("专业技能")
add_skill("CAD/建模：", "SolidWorks（熟练）、Creo（熟练）、UG（熟练）、AutoCAD（熟练）、Moldex 模流分析")
add_skill("编程语言：", "Python（主力）、C/C++（嵌入式）、MATLAB（数据分析与建模）")
add_skill("机器人开发：", "ROS2 框架、SLAM 导航建图、路径规划、传感器融合")
add_skill("嵌入式硬件：", "ESP32 开发、PCB 设计（嘉立创 EDA）、PLC 编程（博途）")
add_skill("AI 工具：", "Claude Code、Codex 等 AI 辅助编程工具，AI Agent 开发经验")
add_skill("仿真分析：", "伏图软件（有限元分析）、MATLAB 信号处理")
add_skill("办公软件：", "Office 全家桶（Word/Excel/PPT 熟练使用）")

# ---- 获奖 ----
add_section("获奖与荣誉")
add_bullet("挑战杯广东省省赛一等奖 — 一鉴钟氢（光声衰荡光谱氢气检测技术）")
add_bullet("GitHub 个人仓库 20+ 个项目，涵盖机器人、嵌入式、数据分析、AI 应用等方向")

# ---- 自我评价 ----
add_section("自我评价")
add_bullet("智能制造工程专业背景，具备机械结构设计 + 嵌入式开发 + AI 应用的复合能力。")
add_bullet("自驱力强，外向开朗，团队协作能力优秀，善于将技术方案落地为可执行的产品。")

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
